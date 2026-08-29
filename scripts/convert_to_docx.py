#!/usr/bin/env python3
"""
Convert the KMS_User_Guide.md to .docx using python-docx.

All screenshots are embedded as images. Handles headings, paragraphs,
bullet lists, ordered lists, tables, and horizontal rules.

Usage: poetry run python scripts/convert_to_docx.py
"""

import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.dml.color import ColorFormat

# ── Paths ──────────────────────────────────────────────────────────────────
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(ROOT, "docs", "KMS_User_Guide.md")
IMG_DIR = os.path.join(ROOT, "docs", "images")
DOCX_PATH = os.path.join(ROOT, "docs", "KMS_User_Guide.docx")

# ── Colours ────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEAL = RGBColor(0x0E, 0x7A, 0x6B)
GREY_DARK = RGBColor(0x33, 0x33, 0x33)
GREY_MID = RGBColor(0x5A, 0x6C, 0x7D)
BLUE_LINK = RGBColor(0x1A, 0x5C, 0xB0)
LIGHT_BG = "F0F4F8"
TABLE_HEADER_BG = "1B2A4A"
TABLE_ALT_BG = "F5F7FA"
ACCENT_TEAL = "0E7A6B"

# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background colour for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size_pt=None, alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    else:
        run.font.size = Pt(8)
    run.font.name = "Calibri"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return run


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="CBD5E0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)


def parse_inline(paragraph, text):
    """Parse bold, italic, inline code, and links inside a paragraph."""
    # Order matters: process links first, then bold, italic, code
    # We'll tokenize and add runs
    # Pattern: **bold**, *italic*, `code`, [text](url)
    token_re = re.compile(
        r"(\[.+?\]\(.+?\)|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)"
    )
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        # Link
        m = re.match(r"\[(.+?)\]\((.+?)\)", part)
        if m:
            label, url = m.group(1), m.group(2)
            run = paragraph.add_run(label)
            run.font.color.rgb = BLUE_LINK
            run.underline = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            continue
        # Inline code
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            rPr = run._element.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
            rPr.append(shd)
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
            continue
        # Bold
        if (part.startswith("**") and part.endswith("**")) or (
            part.startswith("__") and part.endswith("__")
        ):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = GREY_DARK
            continue
        # Italic
        if (part.startswith("*") and part.endswith("*")) or (
            part.startswith("_") and part.endswith("_")
        ):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            continue
        # Plain
        run = paragraph.add_run(part)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = GREY_DARK


def add_styled_paragraph(doc, text, style_name=None, alignment=None, color=None, size_pt=None, bold=False, italic=False, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        parse_inline(p, text)
        # Override size/color if requested
        for run in p.runs:
            if size_pt:
                run.font.size = Pt(size_pt)
            if color:
                run.font.color.rgb = color
            if bold:
                run.bold = True
            if italic:
                run.italic = True
    return p


# ── Table parsing ──────────────────────────────────────────────────────────

def is_table_separator(line):
    stripped = line.strip()
    return bool(re.match(r"^[\|\s\-:]+$", stripped)) and "---" in stripped


def is_table_row(line):
    return "|" in line and line.strip().startswith("|")


def parse_table_row(line):
    # Split on |, strip, drop empty first/last
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def add_table(doc, headers, rows):
    if not headers:
        return
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, TABLE_HEADER_BG)
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=7.5)

    # Data rows
    for i, row in enumerate(rows):
        for j in range(ncols):
            txt = row[j] if j < len(row) else ""
            cell = table.rows[i + 1].cells[j]
            if i % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            # Bold first column
            is_bold = (j == 0)
            set_cell_text(cell, txt, bold=is_bold, size_pt=7.5)

    # Set column widths proportionally (first col a bit wider for labels)
    # Let Word auto-fit; we just set the table width
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Code block ─────────────────────────────────────────────────────────────

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
    pPr.append(shd)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="6" w:color="CBD5E0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="10" w:color="CBD5E0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="6" w:color="CBD5E0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="6" w:color="CBD5E0"/>'
        f'</w:pBdr>'
    )
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for line in lines:
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = GREY_DARK


# ── Image helpers ──────────────────────────────────────────────────────────

def add_image(doc, image_path, caption=None, width_inches=6.2):
    if not os.path.exists(image_path):
        p = doc.add_paragraph()
        run = p.add_run(f"[Image not found: {os.path.basename(image_path)}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    # Add subtle border via paragraph shading trick — picture already has its own
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(1)
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = GREY_MID
        r.font.name = "Calibri"


# ── Cover page ─────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Top accent bar
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT_TEAL}" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(" " * 80)
    run.font.size = Pt(4)
    p.paragraph_format.space_after = Pt(40)

    # Logo placeholder / title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("KMS  STARFARM")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Knowledge Management System")
    run.font.size = Pt(16)
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"
    run.italic = True

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("User Guide  &  Project Overview")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("With screenshots  ·  Step-by-step instructions  ·  KPI framework")
    run.font.size = Pt(10)
    run.font.color.rgb = GREY_MID
    run.font.name = "Calibri"

    # Metadata box
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC" w:val="clear"/>')
    pPr.append(shd)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="8" w:color="CBD5E0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="12" w:color="CBD5E0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="8" w:color="CBD5E0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="12" w:color="CBD5E0"/>'
        f'</w:pBdr>'
    )
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(16)

    for label, value in [
        ("Version", "1.0  —  August 28, 2026"),
        ("System", "KMS Starfarm  v1.0.0"),
        ("Purpose", "Response to CTU deliverables request — project showcase & user manual"),
        ("Status", "Live system  ·  All screenshots taken from the running application"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.3)
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = NAVY
        r1.font.name = "Calibri"
        r2 = p.add_run(value)
        r2.font.size = Pt(9)
        r2.font.color.rgb = GREY_DARK
        r2.font.name = "Calibri"

    # Page break after cover
    doc.add_page_break()


# ── Section divider ────────────────────────────────────────────────────────

def add_section_divider(doc, number, title):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT_TEAL}" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(f"  {number}  —  {title}  ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)


# ── Main conversion ────────────────────────────────────────────────────────

# Map headings to section colours
SECTION_HEADINGS = {
    "1. Introduction": TEAL,
    "2. System Architecture": TEAL,
    "3. Getting Started": TEAL,
    "4. KPI Analysis Dashboard": TEAL,
    "5. Document Management": TEAL,
    "6. AI Chat Assistant (RAG)": TEAL,
    "7. Knowledge Pipeline Configuration": TEAL,
    "8. KPI Extraction & Review Workflow": TEAL,
    "9. Settings & Administration": TEAL,
    "10. Appendix": TEAL,
}

# Image reference pattern: ![alt](images/filename.png)
IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def convert():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

    # ── Default font ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = GREY_DARK
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ── Heading styles ─────────────────────────────────────────────────
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = NAVY if level <= 2 else TEAL
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(20)
            hs.paragraph_format.space_before = Pt(20)
            hs.paragraph_format.space_after = Pt(8)
        elif level == 2:
            hs.font.size = Pt(14)
            hs.paragraph_format.space_before = Pt(14)
            hs.paragraph_format.space_after = Pt(6)
        else:
            hs.font.size = Pt(11)
            hs.paragraph_format.space_before = Pt(8)
            hs.paragraph_format.space_after = Pt(4)

    # ── Title styles ───────────────────────────────────────────────────
    title_style = doc.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = NAVY

    # ── Cover page ─────────────────────────────────────────────────────
    add_cover_page(doc)

    # ── Parse markdown ─────────────────────────────────────────────────
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""
    in_table = False
    table_headers = []
    table_rows = []
    # For bullet list tracking
    list_indent_stack = []

    # Skip lines until after the cover-like content (title already on cover)
    # We'll process everything but style the first H1 specially
    first_h1_done = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code block ──
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Flush pending table ──
        def flush_table():
            nonlocal in_table, table_headers, table_rows
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        # ── Table ──
        if is_table_row(line):
            # Check if next non-empty line is a separator -> this is header
            if not in_table:
                # Peek ahead
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and is_table_separator(lines[j]):
                    table_headers = parse_table_row(line)
                    in_table = True
                    i += 1  # skip header
                    # skip separator line
                    if j < len(lines) and is_table_separator(lines[j]):
                        i = j + 1
                    continue
                else:
                    # No header separator — treat as data row with no header
                    if not in_table:
                        table_headers = []
                        in_table = True
                    table_rows.append(parse_table_row(line))
                    i += 1
                    continue
            else:
                # Already in table — data row (skip separator lines)
                if is_table_separator(line):
                    i += 1
                    continue
                table_rows.append(parse_table_row(line))
                i += 1
                continue
        else:
            # Not a table row — flush if we were in a table
            if in_table:
                flush_table()

        # ── Horizontal rule ──
        if stripped in ("---", "***", "___") or re.match(r"^-{3,}$", stripped):
            add_horizontal_line(doc)
            i += 1
            continue

        # ── Image ──
        img_match = IMG_RE.search(line)
        if img_match:
            alt_text = img_match.group(1)
            img_rel = img_match.group(2)
            # Resolve relative to docs/
            img_path = os.path.join(os.path.dirname(MD_PATH), img_rel)
            if not os.path.exists(img_path):
                # Try relative to images dir
                img_path = os.path.join(IMG_DIR, os.path.basename(img_rel))
            # Extract any text before/after the image on same line
            before = line[: img_match.start()].strip()
            after = line[img_match.end() :].strip()
            if before:
                add_styled_paragraph(doc, before)
            add_image(doc, img_path, caption=alt_text if alt_text else None)
            if after:
                add_styled_paragraph(doc, after)
            i += 1
            continue

        # ── Heading ──
        hm = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            # Remove markdown heading anchor noise
            text = re.sub(r"\s*\{#.*?\}\s*$", "", text)

            # Skip the very first H1 (it's the cover page title, already rendered)
            if level == 1 and not first_h1_done:
                first_h1_done = True
                # Add a subtle title repetition as section start
                # Actually add it as heading 1
                h = doc.add_heading(text, level=1)
                # Add decorative line
                add_horizontal_line(doc)
                i += 1
                continue

            # Check if it's a numbered top-level section heading
            if level == 2 and re.match(r"^\d+\.", text):
                # Use section divider for main sections
                pass  # Keep as normal heading 2

            h = doc.add_heading(text, level=min(level, 3))
            # For H2 with numbering, add accent
            if level == 2 and re.match(r"^\d+\.", text):
                for run in h.runs:
                    run.font.color.rgb = TEAL
            i += 1
            continue

        # ── Blockquote ──
        if stripped.startswith(">"):
            bq_text = re.sub(r"^>\s*", "", stripped)
            # Collect consecutive blockquote lines
            bq_lines = [bq_text]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith(">"):
                bq_lines.append(re.sub(r"^>\s*", "", lines[j].strip()))
                j += 1
            full_bq = " ".join(bq_lines)
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            # Left border
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{ACCENT_TEAL}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0FDF9" w:val="clear"/>')
            pPr.append(shd)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Strip markdown formatting inside blockquote for simplicity
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", full_bq)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            run = p.add_run(clean)
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = GREY_MID
            run.font.name = "Calibri"
            i = j
            continue

        # ── Bullet / ordered list ──
        bullet_m = re.match(r"^(\s*)([-*+])\s+(.*)", line)
        ordered_m = re.match(r"^(\s*)(\d+)\.\s+(.*)", line)
        if bullet_m or ordered_m:
            is_ordered = bool(ordered_m)
            if is_ordered:
                indent = len(ordered_m.group(1))
                num = ordered_m.group(2)
                content = ordered_m.group(3)
            else:
                indent = len(bullet_m.group(1))
                content = bullet_m.group(3)

            level = indent // 2  # 2 spaces per level
            p = doc.add_paragraph(style="List Bullet" if not is_ordered else "List Number")
            if level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            # Remove bold markers for list content and parse inline
            parse_inline(p, content)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"
                if not run.bold:
                    run.font.color.rgb = GREY_DARK
            i += 1
            continue

        # ── Empty line ──
        if not stripped:
            i += 1
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        parse_inline(p, line)
        for run in p.runs:
            if run.font.size is None or run.font.size.pt is None:
                run.font.size = Pt(10)
        i += 1

    # Flush any remaining table
    if in_table and table_headers:
        add_table(doc, table_headers, table_rows)

    # ── Footer ─────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("KMS Starfarm  ·  Knowledge Management System  ·  Confidential")
        run.font.size = Pt(7)
        run.font.color.rgb = GREY_MID
        run.font.name = "Calibri"
        run.italic = True

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("KMS Starfarm  —  User Guide  v1.0")
        run.font.size = Pt(7)
        run.font.color.rgb = GREY_MID
        run.font.name = "Calibri"

    # ── Save ───────────────────────────────────────────────────────────
    doc.save(DOCX_PATH)
    size_kb = os.path.getsize(DOCX_PATH) / 1024
    print(f"Saved: {DOCX_PATH}  ({size_kb:.0f} KB)")


if __name__ == "__main__":
    convert()
